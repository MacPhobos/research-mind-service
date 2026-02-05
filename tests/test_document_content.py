"""Integration tests for document content type API endpoints.

Tests cover:
- Document uploads (PDF, DOCX, TXT, MD)
- Content extraction and storage
- Title override
- Metadata in response
- Error handling (unsupported format, encrypted PDF, empty file, corrupted file)

Requires:
- fitz (PyMuPDF) for creating test PDFs
- python-docx for creating test DOCX files
"""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import patch

import fitz  # PyMuPDF
import pytest
from docx import Document  # python-docx
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401  -- ensure models registered with Base.metadata
from app.db.base import Base
from app.db.session import get_db
from app.main import app


# ------------------------------------------------------------------
# Fixtures
# ------------------------------------------------------------------


@pytest.fixture()
def tmp_content_sandbox(tmp_path: Path) -> str:
    """Provide a temporary content sandbox root directory."""
    return str(tmp_path / "content_sandboxes")


@pytest.fixture()
def db_engine():
    """Create an in-memory SQLite engine shared across connections.

    Enables foreign key enforcement for proper CASCADE behavior testing.
    """
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    # Enable foreign key enforcement in SQLite
    from sqlalchemy import event

    @event.listens_for(engine, "connect")
    def set_sqlite_pragma(dbapi_connection, connection_record):
        cursor = dbapi_connection.cursor()
        cursor.execute("PRAGMA foreign_keys=ON")
        cursor.close()

    Base.metadata.create_all(bind=engine)
    yield engine
    Base.metadata.drop_all(bind=engine)
    engine.dispose()


@pytest.fixture()
def db_session(db_engine) -> Session:
    """Yield a SQLAlchemy session bound to the shared in-memory engine."""
    TestingSessionLocal = sessionmaker(
        autocommit=False, autoflush=False, bind=db_engine
    )
    session = TestingSessionLocal()
    try:
        yield session
    finally:
        session.rollback()
        session.close()


@pytest.fixture()
def client(db_engine, tmp_content_sandbox: str) -> TestClient:
    """TestClient with overridden DB and content_sandbox_root."""
    from app.core.config import settings

    # Save original settings
    original_content_sandbox_root = settings.content_sandbox_root

    # Override settings (bypass pydantic model immutability)
    object.__setattr__(settings, "content_sandbox_root", tmp_content_sandbox)

    TestingSessionLocal = sessionmaker(
        autocommit=False, autoflush=False, bind=db_engine
    )

    def _override_get_db():
        session = TestingSessionLocal()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db] = _override_get_db
    with TestClient(app) as tc:
        yield tc
    app.dependency_overrides.clear()

    # Restore original settings
    object.__setattr__(settings, "content_sandbox_root", original_content_sandbox_root)


def _create_session(client: TestClient, name: str = "Test Session") -> dict:
    """Helper to create a session and return the response data."""
    response = client.post(
        "/api/v1/sessions/",
        json={"name": name, "description": "Test session for document tests"},
    )
    assert response.status_code == 201, f"Failed to create session: {response.text}"
    return response.json()


# ------------------------------------------------------------------
# Test Document Helpers
# ------------------------------------------------------------------


def create_test_pdf(path: Path, content: str = "Test PDF Content") -> None:
    """Create a simple PDF document with text content.

    Args:
        path: Path to write the PDF file.
        content: Text content to include in the PDF.
    """
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), content, fontsize=12)
    doc.save(str(path))
    doc.close()


def create_encrypted_pdf(path: Path, password: str = "secret123") -> None:
    """Create an encrypted PDF document.

    Args:
        path: Path to write the PDF file.
        password: Password for encryption.
    """
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "This is encrypted content", fontsize=12)
    # Save with encryption
    doc.save(
        str(path),
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw=password,
        owner_pw=password,
    )
    doc.close()


def create_test_docx(path: Path, content: str = "Test DOCX Content") -> None:
    """Create a simple DOCX document with text content.

    Args:
        path: Path to write the DOCX file.
        content: Text content to include in the document.
    """
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph(content)
    doc.save(str(path))


def create_empty_file(path: Path) -> None:
    """Create an empty file.

    Args:
        path: Path to the file.
    """
    path.write_text("")


def create_corrupted_docx(path: Path) -> None:
    """Create a corrupted DOCX file (random bytes with .docx extension).

    Args:
        path: Path to write the corrupted file.
    """
    import random

    # Write random bytes that are not valid DOCX structure
    random_bytes = bytes(random.getrandbits(8) for _ in range(500))
    path.write_bytes(random_bytes)


# ------------------------------------------------------------------
# Success Case Tests
# ------------------------------------------------------------------


class TestDocumentContentSuccess:
    """Tests for successful document content operations."""

    def test_add_document_pdf_success(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """POST PDF document returns 201 with status=ready and extracted content."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test PDF
        pdf_path = tmp_path / "test_document.pdf"
        create_test_pdf(pdf_path, "This is the PDF test content for extraction.")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "My PDF Document",
                "source": str(pdf_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["session_id"] == session_id
        assert data["content_type"] == "document"
        assert data["title"] == "My PDF Document"
        assert data["status"] == "ready"
        assert data["size_bytes"] > 0
        assert data["mime_type"] == "text/markdown"
        assert "content_id" in data
        assert "created_at" in data

        # Verify extracted content was written to sandbox
        content_dir = Path(tmp_content_sandbox) / session_id / data["content_id"]
        assert content_dir.exists()
        content_file = content_dir / "content.md"
        assert content_file.exists()

        extracted_content = content_file.read_text()
        assert "PDF test content" in extracted_content

        # Verify metadata.json was created
        metadata_file = content_dir / "metadata.json"
        assert metadata_file.exists()
        metadata = json.loads(metadata_file.read_text())
        assert metadata["original_filename"] == "test_document.pdf"
        assert metadata["file_extension"] == ".pdf"
        assert "extraction_method" in metadata

    def test_add_document_docx_success(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """POST DOCX document returns 201 with markdown conversion."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test DOCX
        docx_path = tmp_path / "test_document.docx"
        create_test_docx(docx_path, "This is the DOCX test content for extraction.")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "My Word Document",
                "source": str(docx_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["session_id"] == session_id
        assert data["content_type"] == "document"
        assert data["title"] == "My Word Document"
        assert data["status"] == "ready"
        assert data["size_bytes"] > 0
        assert data["mime_type"] == "text/markdown"

        # Verify content file is markdown (content.md)
        content_dir = Path(tmp_content_sandbox) / session_id / data["content_id"]
        content_file = content_dir / "content.md"
        assert content_file.exists()

        extracted_content = content_file.read_text()
        # Heading should be converted to markdown
        assert "Test Document" in extracted_content
        assert "DOCX test content" in extracted_content

    def test_add_document_txt_success(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """POST TXT document returns 201 with content.txt created."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test TXT file
        txt_path = tmp_path / "test_document.txt"
        txt_path.write_text("This is plain text content for testing.", encoding="utf-8")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "My Text File",
                "source": str(txt_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["session_id"] == session_id
        assert data["content_type"] == "document"
        assert data["title"] == "My Text File"
        assert data["status"] == "ready"
        assert data["mime_type"] == "text/plain"

        # Verify content.txt was created (not content.md for .txt files)
        content_dir = Path(tmp_content_sandbox) / session_id / data["content_id"]
        content_file = content_dir / "content.txt"
        assert content_file.exists()
        assert content_file.read_text() == "This is plain text content for testing."

    def test_add_document_md_success(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """POST MD document returns 201 with passthrough to content.md."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test MD file
        md_path = tmp_path / "test_document.md"
        md_content = "# Heading\n\nThis is **markdown** content."
        md_path.write_text(md_content, encoding="utf-8")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "My Markdown File",
                "source": str(md_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["session_id"] == session_id
        assert data["content_type"] == "document"
        assert data["title"] == "My Markdown File"
        assert data["status"] == "ready"
        assert data["mime_type"] == "text/markdown"

        # Verify content.md was created with passthrough content
        content_dir = Path(tmp_content_sandbox) / session_id / data["content_id"]
        content_file = content_dir / "content.md"
        assert content_file.exists()
        assert content_file.read_text() == md_content

    def test_add_document_with_title_override(
        self, client: TestClient, tmp_path: Path
    ):
        """POST document with custom title uses that title instead of filename."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test PDF with generic filename
        pdf_path = tmp_path / "file123.pdf"
        create_test_pdf(pdf_path, "Some content")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Custom Document Title",
                "source": str(pdf_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["title"] == "Custom Document Title"

    def test_add_document_metadata_returned(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """POST document returns extraction metadata in response."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test PDF
        pdf_path = tmp_path / "metadata_test.pdf"
        create_test_pdf(pdf_path, "Content for metadata test")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Metadata Test",
                "source": str(pdf_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()

        # Verify metadata_json contains extraction details
        assert data["metadata_json"] is not None
        metadata = data["metadata_json"]

        # Should have document-specific metadata
        assert "original_filename" in metadata
        assert metadata["original_filename"] == "metadata_test.pdf"
        assert "file_extension" in metadata
        assert metadata["file_extension"] == ".pdf"
        assert "extraction_method" in metadata
        assert "extracted_at" in metadata
        assert "content_stats" in metadata

        # Content stats should have word/char counts
        content_stats = metadata["content_stats"]
        assert "word_count" in content_stats
        assert "char_count" in content_stats
        assert content_stats["word_count"] > 0

    def test_add_document_without_title_uses_filename(
        self, client: TestClient, tmp_path: Path
    ):
        """POST document without title uses file stem as default title."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test PDF with descriptive filename
        pdf_path = tmp_path / "important_research_notes.pdf"
        create_test_pdf(pdf_path, "Research content")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "source": str(pdf_path),
            },
        )

        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        # Title should be the file stem (without extension)
        assert data["title"] == "important_research_notes"


# ------------------------------------------------------------------
# Error Case Tests
# ------------------------------------------------------------------


class TestDocumentContentErrors:
    """Tests for document content error handling."""

    def test_add_document_unsupported_format_error(
        self, client: TestClient, tmp_path: Path
    ):
        """POST document with unsupported format returns status=error."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create a file with unsupported extension
        xlsx_path = tmp_path / "spreadsheet.xlsx"
        xlsx_path.write_bytes(b"fake xlsx content")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Spreadsheet",
                "source": str(xlsx_path),
            },
        )

        # API returns 201 but with status=error
        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["status"] == "error"
        assert data["error_message"] is not None
        assert "unsupported" in data["error_message"].lower()
        assert ".xlsx" in data["error_message"]

    def test_add_document_encrypted_pdf_error(
        self, client: TestClient, tmp_path: Path
    ):
        """POST encrypted PDF returns status=error."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create encrypted PDF
        pdf_path = tmp_path / "encrypted.pdf"
        create_encrypted_pdf(pdf_path, password="secret123")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Encrypted PDF",
                "source": str(pdf_path),
            },
        )

        # API returns 201 but with status=error
        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["status"] == "error"
        assert data["error_message"] is not None
        assert "encrypted" in data["error_message"].lower()

    def test_add_document_empty_file_error(
        self, client: TestClient, tmp_path: Path
    ):
        """POST empty TXT file returns status=error."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create empty text file
        txt_path = tmp_path / "empty.txt"
        create_empty_file(txt_path)

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Empty File",
                "source": str(txt_path),
            },
        )

        # API returns 201 but with status=error
        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["status"] == "error"
        assert data["error_message"] is not None
        assert "empty" in data["error_message"].lower()

    def test_add_document_corrupted_file_error(
        self, client: TestClient, tmp_path: Path
    ):
        """POST corrupted DOCX file returns status=error."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create corrupted DOCX
        docx_path = tmp_path / "corrupted.docx"
        create_corrupted_docx(docx_path)

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Corrupted DOCX",
                "source": str(docx_path),
            },
        )

        # API returns 201 but with status=error
        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["status"] == "error"
        assert data["error_message"] is not None
        assert "corrupt" in data["error_message"].lower() or "failed" in data["error_message"].lower()

    def test_add_document_file_not_found_error(
        self, client: TestClient, tmp_path: Path
    ):
        """POST with non-existent file path returns status=error."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Use a path that doesn't exist
        nonexistent_path = tmp_path / "does_not_exist.pdf"

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Missing File",
                "source": str(nonexistent_path),
            },
        )

        # API returns 201 but with status=error
        assert response.status_code == 201, f"Response: {response.text}"
        data = response.json()
        assert data["status"] == "error"
        assert data["error_message"] is not None
        assert "not found" in data["error_message"].lower()

    def test_add_document_invalid_session(
        self, client: TestClient, tmp_path: Path
    ):
        """POST to non-existent session returns 404."""
        fake_session_id = "00000000-0000-4000-a000-000000000000"

        # Create a valid document
        pdf_path = tmp_path / "test.pdf"
        create_test_pdf(pdf_path, "Test content")

        response = client.post(
            f"/api/v1/sessions/{fake_session_id}/content/",
            data={
                "content_type": "document",
                "title": "Test",
                "source": str(pdf_path),
            },
        )

        assert response.status_code == 404
        data = response.json()
        assert data["detail"]["error"]["code"] == "SESSION_NOT_FOUND"


# ------------------------------------------------------------------
# Content Lifecycle Tests
# ------------------------------------------------------------------


class TestDocumentContentLifecycle:
    """Tests for document content item lifecycle operations."""

    def test_document_content_can_be_deleted(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """DELETE removes document content item and storage files."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create and add document
        pdf_path = tmp_path / "to_delete.pdf"
        create_test_pdf(pdf_path, "Content to be deleted")

        create_response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Delete Me",
                "source": str(pdf_path),
            },
        )
        assert create_response.status_code == 201
        content_id = create_response.json()["content_id"]

        # Verify storage exists
        content_dir = Path(tmp_content_sandbox) / session_id / content_id
        assert content_dir.exists()

        # Delete content
        delete_response = client.delete(
            f"/api/v1/sessions/{session_id}/content/{content_id}"
        )
        assert delete_response.status_code == 204

        # Verify storage is removed
        assert not content_dir.exists()

        # Verify GET returns 404
        get_response = client.get(
            f"/api/v1/sessions/{session_id}/content/{content_id}"
        )
        assert get_response.status_code == 404

    def test_document_appears_in_content_list(
        self, client: TestClient, tmp_path: Path
    ):
        """Document content items appear in session content list."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Add multiple document types
        pdf_path = tmp_path / "doc.pdf"
        txt_path = tmp_path / "doc.txt"
        create_test_pdf(pdf_path, "PDF content")
        txt_path.write_text("TXT content", encoding="utf-8")

        client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={"content_type": "document", "title": "PDF Doc", "source": str(pdf_path)},
        )
        client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={"content_type": "document", "title": "TXT Doc", "source": str(txt_path)},
        )

        # List content
        response = client.get(f"/api/v1/sessions/{session_id}/content/")
        assert response.status_code == 200
        data = response.json()

        assert data["count"] == 2
        content_types = [item["content_type"] for item in data["items"]]
        assert all(ct == "document" for ct in content_types)

        titles = [item["title"] for item in data["items"]]
        assert "PDF Doc" in titles
        assert "TXT Doc" in titles

    def test_document_original_file_not_retained(
        self, client: TestClient, tmp_content_sandbox: str, tmp_path: Path
    ):
        """Original document file is NOT copied to storage (only extracted content)."""
        session = _create_session(client)
        session_id = session["session_id"]

        # Create test PDF
        pdf_path = tmp_path / "original.pdf"
        create_test_pdf(pdf_path, "Original PDF content")

        response = client.post(
            f"/api/v1/sessions/{session_id}/content/",
            data={
                "content_type": "document",
                "title": "Test Original Not Retained",
                "source": str(pdf_path),
            },
        )
        assert response.status_code == 201
        data = response.json()
        content_id = data["content_id"]

        # Check storage directory
        content_dir = Path(tmp_content_sandbox) / session_id / content_id
        files_in_dir = list(content_dir.iterdir())
        filenames = [f.name for f in files_in_dir]

        # Should have content.md and metadata.json, but NOT original.pdf
        assert "content.md" in filenames
        assert "metadata.json" in filenames
        assert "original.pdf" not in filenames
        assert not any(f.suffix == ".pdf" for f in files_in_dir)
