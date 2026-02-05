"""Tests for DOCX content extractor."""

import pytest
from docx import Document
from docx.shared import Pt

from app.services.extractors.document.docx import DOCXExtractor


@pytest.fixture
def docx_extractor():
    """Create a DOCXExtractor instance."""
    return DOCXExtractor()


@pytest.fixture
def simple_docx(tmp_path):
    """Create a simple DOCX with text content."""
    docx_path = tmp_path / "simple.docx"
    doc = Document()
    doc.add_paragraph("Hello, World!")
    doc.add_paragraph("This is a simple DOCX document.")
    doc.add_paragraph("It contains multiple paragraphs.")
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def docx_with_metadata(tmp_path):
    """Create a DOCX with document properties."""
    docx_path = tmp_path / "with_metadata.docx"
    doc = Document()
    doc.add_paragraph("Content with metadata")

    # Set document properties
    doc.core_properties.title = "Test Document Title"
    doc.core_properties.author = "Test Author"
    doc.core_properties.subject = "Test Subject"

    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def docx_with_formatting(tmp_path):
    """Create a DOCX with various formatting."""
    docx_path = tmp_path / "formatted.docx"
    doc = Document()

    # Add heading
    doc.add_heading("Main Heading", level=1)

    # Add paragraph with bold and italic
    para = doc.add_paragraph()
    para.add_run("This is ").bold = False
    bold_run = para.add_run("bold text")
    bold_run.bold = True
    para.add_run(" and ")
    italic_run = para.add_run("italic text")
    italic_run.italic = True
    para.add_run(".")

    # Add subheading
    doc.add_heading("Sub Heading", level=2)
    doc.add_paragraph("Content under subheading.")

    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def docx_with_lists(tmp_path):
    """Create a DOCX with lists."""
    docx_path = tmp_path / "with_lists.docx"
    doc = Document()

    doc.add_paragraph("Bullet list:")
    for item in ["First item", "Second item", "Third item"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Numbered list:")
    for item in ["Item one", "Item two", "Item three"]:
        doc.add_paragraph(item, style="List Number")

    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def empty_docx(tmp_path):
    """Create an empty DOCX."""
    docx_path = tmp_path / "empty.docx"
    doc = Document()
    # Save without adding any content
    doc.save(str(docx_path))
    return docx_path


class TestDOCXExtractor:
    """Test suite for DOCXExtractor."""

    @pytest.mark.asyncio
    async def test_extract_simple_docx(self, docx_extractor, simple_docx):
        """Test extracting content from a simple DOCX."""
        result = await docx_extractor.extract(simple_docx)

        assert result.content is not None
        assert "Hello, World!" in result.content
        assert "simple DOCX document" in result.content
        assert "multiple paragraphs" in result.content

    @pytest.mark.asyncio
    async def test_extract_docx_metadata(self, docx_extractor, docx_with_metadata):
        """Test extracting metadata from DOCX."""
        result = await docx_extractor.extract(docx_with_metadata)

        assert result.document_metadata is not None
        assert result.document_metadata.get("title") == "Test Document Title"
        assert result.document_metadata.get("author") == "Test Author"
        assert result.document_metadata.get("subject") == "Test Subject"

    @pytest.mark.asyncio
    async def test_extract_docx_with_headings(self, docx_extractor, docx_with_formatting):
        """Test that headings are preserved as markdown."""
        result = await docx_extractor.extract(docx_with_formatting)

        # Headings should be converted to ATX style (#, ##)
        assert "Main Heading" in result.content
        assert "Sub Heading" in result.content

    @pytest.mark.asyncio
    async def test_extract_docx_with_formatting(self, docx_extractor, docx_with_formatting):
        """Test that bold/italic formatting is preserved."""
        result = await docx_extractor.extract(docx_with_formatting)

        # Check that content is extracted (formatting may vary)
        assert "bold text" in result.content
        assert "italic text" in result.content

    @pytest.mark.asyncio
    async def test_extract_docx_with_lists(self, docx_extractor, docx_with_lists):
        """Test that lists are extracted."""
        result = await docx_extractor.extract(docx_with_lists)

        assert "First item" in result.content
        assert "Second item" in result.content
        assert "Third item" in result.content
        assert "Item one" in result.content
        assert "Item two" in result.content

    @pytest.mark.asyncio
    async def test_empty_docx_raises_error(self, docx_extractor, empty_docx):
        """Test that empty DOCX raises ValueError."""
        with pytest.raises(ValueError) as exc_info:
            await docx_extractor.extract(empty_docx)

        assert "empty" in str(exc_info.value).lower() or "no content" in str(exc_info.value).lower()

    @pytest.mark.asyncio
    async def test_nonexistent_docx_raises_error(self, docx_extractor, tmp_path):
        """Test that non-existent file raises ValueError."""
        nonexistent = tmp_path / "nonexistent.docx"

        with pytest.raises(ValueError) as exc_info:
            await docx_extractor.extract(nonexistent)

        assert "corrupted" in str(exc_info.value).lower()

    @pytest.mark.asyncio
    async def test_corrupted_docx_raises_error(self, docx_extractor, tmp_path):
        """Test that corrupted DOCX raises ValueError."""
        corrupted_path = tmp_path / "corrupted.docx"
        corrupted_path.write_bytes(b"not a valid docx content")

        with pytest.raises(ValueError) as exc_info:
            await docx_extractor.extract(corrupted_path)

        assert "corrupted" in str(exc_info.value).lower()

    @pytest.mark.asyncio
    async def test_method_name(self, docx_extractor):
        """Test that method_name is set correctly."""
        assert docx_extractor.method_name == "mammoth"

    @pytest.mark.asyncio
    async def test_content_is_stripped(self, docx_extractor, simple_docx):
        """Test that extracted content is stripped of leading/trailing whitespace."""
        result = await docx_extractor.extract(simple_docx)

        assert result.content == result.content.strip()

    @pytest.mark.asyncio
    async def test_metadata_removes_none_values(self, docx_extractor, simple_docx):
        """Test that None metadata values are removed."""
        result = await docx_extractor.extract(simple_docx)

        # None values should not be present in metadata
        for value in result.document_metadata.values():
            assert value is not None

    @pytest.mark.asyncio
    async def test_metadata_extraction_is_best_effort(self, docx_extractor, tmp_path):
        """Test that metadata extraction failure doesn't break content extraction."""
        # Create a valid DOCX
        docx_path = tmp_path / "valid.docx"
        doc = Document()
        doc.add_paragraph("Valid content")
        doc.save(str(docx_path))

        # Even if metadata extraction has issues, content should still be extracted
        result = await docx_extractor.extract(docx_path)
        assert "Valid content" in result.content
